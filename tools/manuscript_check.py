#!/usr/bin/env python3
"""投稿前自检工具(manuscript pre-submission checker)。

给一篇英文论文(.txt/.md/.tex),自动跑一遍常见 desk-reject 与语言风险检查,
输出一份分项报告。纯标准库、零依赖、零成本。

用法:
    python tools/manuscript_check.py paper.txt
    python tools/manuscript_check.py paper.tex --target-words 6000 --out report.md
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

# ============ 规则词表 ============
WORDY = {
    "in order to": "to",
    "due to the fact that": "because",
    "a large number of": "many",
    "a number of": "several",
    "it is worth noting that": "(直接陈述重点)",
    "it should be noted that": "(直接陈述重点)",
    "as is well known": "(删,或补引用)",
    "studied and investigated": "investigated",
    "the influence and effect": "the effect",
    "with the help of": "using",
    "in the process of": "during",
    "make use of": "use",
    "in spite of the fact that": "although",
    "has the ability to": "can",
}
VAGUE = ["very", "really", "quite", "pretty", "a lot", "good", "bad", "big", "huge", "nice"]
SUBJECTIVE = ["we think", "we feel", "we believe", "i think", "obviously", "of course", "everyone knows"]
DECLARATIONS = {
    "伦理审批 (ethics approval)": [r"ethic", r"institutional review board", r"\bIRB\b"],
    "知情同意 (informed consent)": [r"informed consent"],
    "利益冲突 (competing interests)": [r"conflict[s]? of interest", r"competing interest"],
    "数据可得性 (data availability)": [r"data availability", r"data .{0,20}available", r"data can be found"],
    "资助声明 (funding)": [r"funding", r"funded by", r"\bgrant\b"],
}
KEY_SECTIONS = ["abstract", "introduction", "method", "result", "discussion", "conclusion", "reference"]

PASSIVE_RE = re.compile(
    r"\b(is|are|was|were|be|been|being)\s+(\w+ed|done|made|shown|given|found|used|proposed|"
    r"obtained|performed|conducted|observed|reported|presented|described|developed|considered)\b",
    re.I,
)
WORD_RE = re.compile(r"[A-Za-z][A-Za-z'\-]*")


# ============ 文本预处理 ============
def read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise SystemExit("读取 .docx 需要 python-docx:pip install python-docx")
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def strip_latex(text: str) -> str:
    text = re.sub(r"(?m)%.*$", "", text)                                  # 注释
    text = re.sub(r"\\(cite|ref|label|eqref)\{[^}]*\}", " [ref] ", text)  # 引用类
    text = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^}]*\})?", " ", text)  # 其它命令
    return text.replace("{", " ").replace("}", " ")


def words(text: str) -> list[str]:
    return WORD_RE.findall(text)


def split_sentences(text: str) -> list[str]:
    # 先按行切(标题/章节名各自独立,避免无标点的标题并进正文句),再按句末标点切
    out: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        for s in re.split(r"(?<=[.!?])\s+", line):
            s = s.strip()
            if len(s.split()) >= 3:
                out.append(s)
    return out


# ============ 检查项:返回 (level, title, findings) ============
# level: "ok" | "warn" | "risk"
def check_word_count(plain: str, target: int | None) -> tuple:
    n = len(words(plain))
    level = "ok"
    findings = [f"全文约 **{n}** 词"]
    if target:
        if n > target:
            level = "warn"
            findings.append(f"超过目标上限 {target} 词,需精简约 {n - target} 词")
        else:
            findings.append(f"在目标上限 {target} 词以内 ✓")
    m = re.search(r"(?is)abstract(.{50,2000}?)(introduction|keywords)", plain)
    if m:
        an = len(words(m.group(1)))
        findings.append(f"摘要约 **{an}** 词" + ("(多数刊 ≤250,偏长)" if an > 250 else ""))
        if an > 300:
            level = "warn"
    return level, "字数 / 长度", findings


def check_sections(plain: str) -> tuple:
    low = plain.lower()
    found = [s for s in KEY_SECTIONS if s in low]
    missing = [s for s in ["method", "result", "discussion", "reference"] if s not in low]
    findings = [f"检测到结构关键词:{', '.join(found) if found else '几乎没有'}"]
    level = "ok"
    if missing:
        level = "warn"
        findings.append(f"未检测到:{', '.join(missing)}(确认是否缺失或用了别的标题)")
    return level, "章节结构", findings


def check_long_sentences(sentences: list[str]) -> tuple:
    longs = sorted(((len(s.split()), s) for s in sentences if len(s.split()) > 40), reverse=True)
    findings = [f"超长句(>40 词)共 **{len(longs)}** 句"]
    level = "ok" if not longs else ("warn" if len(longs) <= 8 else "risk")
    for n, s in longs[:3]:
        snippet = s[:120] + ("…" if len(s) > 120 else "")
        findings.append(f'- ({n} 词)"{snippet}"')
    if longs:
        findings.append("建议拆成 2–3 个短句,审稿人更易读。")
    return level, "超长句", findings


def check_passive(sentences: list[str]) -> tuple:
    if not sentences:
        return "ok", "被动语态", ["无正文句子"]
    hits = sum(1 for s in sentences if PASSIVE_RE.search(s))
    ratio = hits / len(sentences)
    level = "warn" if ratio > 0.45 else "ok"
    findings = [f"含被动结构的句子约占 **{ratio:.0%}**(粗略估计)"]
    if level == "warn":
        findings.append("被动偏多,适当改主动(We …)更有力、更易读。")
    return level, "被动语态", findings


def check_wordy(plain: str) -> tuple:
    low = plain.lower()
    hits = [(p, repl) for p, repl in WORDY.items() if p in low]
    level = "ok" if not hits else "warn"
    findings = [f'- "{p}" → 建议 **{repl}**' for p, repl in hits] or ["未发现常见冗余/中式表达 ✓"]
    return level, "冗余 / 中式表达", findings


def check_vague_subjective(plain: str) -> tuple:
    low = plain.lower()
    v = [w for w in VAGUE if re.search(r"\b" + re.escape(w) + r"\b", low)]
    s = [w for w in SUBJECTIVE if w in low]
    findings = []
    if v:
        findings.append("模糊词:" + ", ".join(f'"{w}"' for w in v) + " → 用具体、可量化的表述")
    if s:
        findings.append("主观表述:" + ", ".join(f'"{w}"' for w in s) + " → 改为基于证据(results suggest …)")
    level = "ok" if not findings else "warn"
    return level, "模糊 / 主观措辞", (findings or ["未发现明显模糊/主观措辞 ✓"])


def check_references(text: str) -> tuple:
    m = re.search(r"(?is)\breferences?\b", text)
    if not m:
        return "warn", "参考文献", ["未找到 References 段落,确认是否遗漏。"]
    tail = text[m.end():]
    bracket = len(re.findall(r"\[\d+\]", tail))
    numbered = len(re.findall(r"(?m)^\s*\d+\.\s", tail))
    dois = len(re.findall(r"10\.\d{4,9}/\S+", text))
    bare_urls = len(re.findall(r"https?://", tail))
    count = max(bracket, numbered, dois)
    level = "warn" if bare_urls > 3 else "ok"
    findings = [f"参考文献区估算条目:**{count or '未识别'}**;DOI {dois} 个;裸 URL {bare_urls} 个"]
    if level == "warn":
        findings.append("裸 URL 偏多,多数刊要求规范引用格式而非直接贴链接。")
    return level, "参考文献", findings


def check_declarations(plain: str) -> tuple:
    low = plain.lower()
    findings, missing = [], 0
    for name, pats in DECLARATIONS.items():
        ok = any(re.search(p, low) for p in pats)
        findings.append(("- ✅ " if ok else "- 🔴 缺 ") + name)
        missing += 0 if ok else 1
    level = "ok" if missing == 0 else ("warn" if missing <= 2 else "risk")
    findings.append("(按目标刊要求补齐;并非每刊都强制全部)")
    return level, "声明 / 合规段落", findings


def check_figures_tables(plain: str) -> tuple:
    figs = sorted({int(n) for n in re.findall(r"(?i)\bfig(?:ure)?\.?\s*(\d+)", plain)})
    tabs = sorted({int(n) for n in re.findall(r"(?i)\btable\s*(\d+)", plain)})
    findings = [f"正文引用 图 {figs or '无'};表 {tabs or '无'}"]
    level = "ok"
    for label, nums in (("图", figs), ("表", tabs)):
        if nums:
            gaps = [i for i in range(1, max(nums) + 1) if i not in nums]
            if gaps:
                level = "warn"
                findings.append(f"{label}编号不连续,缺 {gaps}(确认是否漏引用或跳号)")
    return level, "图表编号", findings


# ============ 报告 ============
EMOJI = {"ok": "🟢", "warn": "🟡", "risk": "🔴"}


def render_report(name: str, results: list[tuple]) -> str:
    lines = [f"# 投稿前自检报告 — {name}", "", "| 项目 | 状态 |", "|---|---|"]
    lines += [f"| {title} | {EMOJI[level]} |" for level, title, _ in results]
    lines.append("")
    for level, title, findings in results:
        lines.append(f"## {EMOJI[level]} {title}")
        lines += [f if f.startswith("-") else f"- {f}" for f in findings]
        lines.append("")
    risks = sum(1 for l, _, _ in results if l == "risk")
    warns = sum(1 for l, _, _ in results if l == "warn")
    lines.append("---")
    lines.append(
        f"**小结**:🔴 {risks} 项需重点处理,🟡 {warns} 项建议优化。"
        "本工具只做机器初筛,最终仍需人工(尤其学科判断与审稿回复)把关——这正是真人服务的价值。"
    )
    return "\n".join(lines)


def analyze(raw: str, is_tex: bool = False, target_words: int | None = None) -> list[tuple]:
    """核心分析入口:输入原文,返回 [(level, title, findings), ...]。CLI 与 Web 共用。"""
    plain = strip_latex(raw) if is_tex else raw
    sentences = split_sentences(plain)
    return [
        check_word_count(plain, target_words),
        check_sections(plain),
        check_long_sentences(sentences),
        check_passive(sentences),
        check_wordy(plain),
        check_vague_subjective(plain),
        check_references(raw),
        check_declarations(plain),
        check_figures_tables(plain),
    ]


def main() -> None:
    ap = argparse.ArgumentParser(description="投稿前自检工具(manuscript pre-submission checker)")
    ap.add_argument("path", help="论文文件(.txt/.md/.tex/.docx)")
    ap.add_argument("--target-words", type=int, default=None, help="目标期刊正文词数上限")
    ap.add_argument("--out", default=None, help="把报告写入文件(默认打印到屏幕)")
    args = ap.parse_args()

    path = Path(args.path)
    raw = read_text(path)
    results = analyze(raw, is_tex=path.suffix.lower() == ".tex", target_words=args.target_words)
    report = render_report(path.name, results)
    if args.out:
        Path(args.out).write_text(report, encoding="utf-8")
        print(f"报告已写入 {args.out}")
    else:
        print(report)


if __name__ == "__main__":
    main()
